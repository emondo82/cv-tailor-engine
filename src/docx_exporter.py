from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def set_document_style(doc: Document):
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(42)
        section.right_margin = Pt(42)


def add_name_header(doc: Document, candidate_name: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(candidate_name)
    run.bold = True
    run.font.size = Pt(18)
    p.space_after = Pt(3)


def add_contact_line(doc: Document, candidate_profile: dict):
    parts = []

    location = candidate_profile.get("location") or ""
    phone = candidate_profile.get("phone") or ""
    email = candidate_profile.get("email") or ""
    linkedin = candidate_profile.get("linkedin") or ""

    if location:
        parts.append(location)
    if phone:
        parts.append(phone)
    if email:
        parts.append(email)
    if linkedin:
        parts.append(linkedin)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(" | ".join(parts))
    run.font.size = Pt(10)
    p.space_after = Pt(8)


def add_section_heading(doc: Document, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    p.space_before = Pt(6)
    p.space_after = Pt(4)


def add_body_paragraph(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    p.space_before = Pt(0)
    p.space_after = Pt(1)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.space_before = Pt(0)
    p.space_after = Pt(0)
    return p


def format_education_line(item: dict) -> str:
    degree = item.get("degree", "")
    field = item.get("field", "")
    institution = item.get("institution", "")
    start_date = item.get("start_date", "")
    end_date = item.get("end_date", "")

    line = degree
    if field:
        line += f" in {field}"
    if institution:
        line += f" — {institution}"
    if start_date or end_date:
        line += f" ({start_date} - {end_date})"

    return line


def format_certification_line(cert: dict) -> str:
    name = cert.get("name", "")
    provider = cert.get("provider", "")
    date = cert.get("date", "")

    line = name
    if provider:
        line += f" — {provider}"
    if date:
        line += f" ({date})"

    return line


def format_language_line(lang: dict) -> str:
    language = lang.get("language", "")
    if lang.get("level"):
        return f"{language}: {lang['level']}"

    spoken = lang.get("spoken", "")
    written = lang.get("written", "")
    return f"{language}: Spoken {spoken}, Written {written}"


def export_resume_to_docx(
    resume,
    output_path: Path,
    candidate_name: str = "Tamas Kosina",
    candidate_profile: dict | None = None,
):
    candidate_profile = candidate_profile or {}

    doc = Document()
    set_document_style(doc)

    add_name_header(doc, candidate_name)
    add_contact_line(doc, candidate_profile)

    add_section_heading(doc, "Professional Summary")
    add_body_paragraph(doc, resume.get("summary", ""))

    add_section_heading(doc, "Core Skills")
    skills = resume.get("skills", [])
    if skills:
        add_body_paragraph(doc, " | ".join(skills))
    else:
        add_body_paragraph(doc, "No skills identified.")

    add_section_heading(doc, "Selected Experience")
    for project in resume.get("selected_projects", []):
        company = project.get("company") or "Unknown"
        role = project.get("role") or "Business Analyst"
        client = project.get("client") or ""
        start = project.get("start_date") or ""
        end = project.get("end_date") or ""

        header = company if not role else f"{company} — {role}"
        add_body_paragraph(doc, header, bold=True)

        sub_parts = []
        if client:
            sub_parts.append(client)
        if start and end:
            sub_parts.append(f"{start} - {end}")
        elif start:
            sub_parts.append(start)

        if sub_parts:
            add_body_paragraph(doc, " | ".join(sub_parts))

        for responsibility in (project.get("responsibilities", []) or [])[:4]:
            add_bullet(doc, responsibility)

    add_section_heading(doc, "Education")
    education = resume.get("education", [])
    if education:
        for item in education:
            add_body_paragraph(doc, format_education_line(item))
    else:
        add_body_paragraph(doc, "No education data available.")

    add_section_heading(doc, "Certifications")
    certifications = resume.get("certifications", [])
    if certifications:
        for cert in certifications:
            add_bullet(doc, format_certification_line(cert))
    else:
        add_body_paragraph(doc, "No certifications available.")

    add_section_heading(doc, "Languages")
    languages = resume.get("languages", [])
    if languages:
        for lang in languages:
            add_bullet(doc, format_language_line(lang))
    else:
        add_body_paragraph(doc, "No language data available.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)