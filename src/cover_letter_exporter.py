from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def set_document_style(doc: Document):
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Pt(42)
        section.bottom_margin = Pt(42)
        section.left_margin = Pt(50)
        section.right_margin = Pt(50)


def add_header(doc: Document, candidate_name: str, candidate_profile: dict):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(candidate_name)
    run.bold = True
    run.font.size = Pt(18)
    p.paragraph_format.space_after = Pt(2)

    parts = []
    location = candidate_profile.get("location") or ""
    phone = candidate_profile.get("phone") or ""
    email = candidate_profile.get("email") or ""
    linkedin = candidate_profile.get("linkedin") or ""

    if location:
        parts.append(location)
    if phone:
        parts.append(phone)
    if email:
        parts.append(email)
    if linkedin:
        parts.append(linkedin)

    if parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(parts))
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(12)


def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    return p


def export_cover_letter_to_docx(
    cover_letter_text: str,
    output_path: Path,
    candidate_name: str = "Tamas Kosina",
    candidate_profile: dict | None = None,
):
    candidate_profile = candidate_profile or {}

    doc = Document()
    set_document_style(doc)
    add_header(doc, candidate_name, candidate_profile)

    paragraphs = [p.strip() for p in cover_letter_text.split("\n\n") if p.strip()]

    for para in paragraphs:
        add_paragraph(doc, para)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)