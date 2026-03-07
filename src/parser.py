from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import List

from docx import Document

from models import (
    CandidateProfile,
    LanguageSkill,
    Certification,
    Education,
    ProjectExperience,
    ResumeKnowledgeBase,
)


def read_docx_text(file_path: Path) -> str:
    doc = Document(file_path)
    parts = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts)


def extract_email(text: str) -> str | None:
    match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    return match.group(0) if match else None


def extract_phone(text: str) -> str | None:
    match = re.search(r"(\+\d{1,3}[\s\-]?\d[\d\s\-]{6,})", text)
    return match.group(1).strip() if match else None


def extract_linkedin(text: str) -> str | None:
    match = re.search(r"(www\.linkedin\.com/in/[^\s]+)", text, re.IGNORECASE)
    return match.group(1) if match else None


def parse_master_resume(master_text: str) -> dict:
    name = "Tamas Kosina" if "TAMAS KOSINA" in master_text.upper() else "Unknown"
    email = extract_email(master_text)
    phone = extract_phone(master_text)
    linkedin = extract_linkedin(master_text)

    location = None
    location_match = re.search(r"Budapest,\s*Hungary", master_text, re.IGNORECASE)
    if location_match:
        location = location_match.group(0)

    summary = ""
    summary_match = re.search(
        r"PROFESSIONAL SUMMARY(.*?)(CORE TECHNICAL SKILLS|WORK EXPERIENCE)",
        master_text,
        re.DOTALL | re.IGNORECASE,
    )
    if summary_match:
        summary = summary_match.group(1).strip()

    return {
        "candidate_profile": CandidateProfile(
            name=name,
            email=email,
            phone=phone,
            linkedin=linkedin,
            location=location,
        ),
        "raw_master_resume_summary": summary,
    }


def extract_languages_from_ec(ec_text: str) -> List[LanguageSkill]:
    languages: List[LanguageSkill] = []

    if "Hungarian" in ec_text:
        languages.append(LanguageSkill(language="Hungarian", level="Native"))

    if "English:" in ec_text:
        languages.append(LanguageSkill(language="English", spoken=5, written=5))
    if "German:" in ec_text:
        languages.append(LanguageSkill(language="German", spoken=3, written=2))
    if "French:" in ec_text:
        languages.append(LanguageSkill(language="French", spoken=3, written=2))

    return languages


def extract_education_from_ec(ec_text: str) -> List[Education]:
    education: List[Education] = []

    if "Master’s in business economics" in ec_text or "Master’s degree or equivalent" in ec_text:
        education.append(
            Education(
                degree="Master's Degree",
                field="Business Economics",
                institution="Szent Istvan University",
                start_date="09/2000",
                end_date="06/2006",
            )
        )

    return education


def extract_certifications_from_ec(ec_text: str) -> List[Certification]:
    known_certs = [
        "Lean Six Sigma Black Belt",
        "Agile Project and Product Management University of Maryland",
        "Product Management University of Maryland",
        "PRINCE 2",
        "Scrum Master",
        "ISQTB",
        "ITIL v4",
    ]

    certifications: List[Certification] = []
    for cert in known_certs:
        if cert.lower() in ec_text.lower():
            certifications.append(Certification(name=cert))

    return certifications


def derive_role(company: str, project_description: str) -> str | None:
    company_lower = (company or "").lower()
    desc_lower = (project_description or "").lower()

    if "johnson controls" in company_lower:
        return "Senior Business Analyst"
    if "aligntech" in company_lower:
        return "Business Analyst"
    if "who" in company_lower or "world health organization" in company_lower:
        return "Business Analyst"
    if "cerner" in company_lower:
        return "Service Cloud Product Owner / Senior Business Analyst"
    if "euipo" in company_lower:
        return "Senior Business Analyst"
    if "credit suisse" in company_lower:
        return "Senior IT Business Analyst"
    if "morgan stanley" in company_lower:
        return "IT Business Analyst"
    if "cognizant" in company_lower:
        return "Senior IT Business Analyst"
    if "wdx" in company_lower:
        return "Business Analyst"
    if "bloomberg" in company_lower:
        return "Technical Support / Integration Analyst"
    if "kbc" in company_lower:
        return "Business Analyst"
    if "capita" in company_lower:
        return "Business Analyst"
    if "citi" in company_lower:
        return "Business Analyst"
    if "exxonmobil" in company_lower:
        return "Business Analyst"
    if "one4all consulting" in company_lower and "capacity4development" in desc_lower:
        return "Product Manager / Senior Business Analyst"
    if "one4all consulting" in company_lower:
        return "Business Analyst"

    return None


def clean_date_text(date_text: str | None) -> str:
    if not date_text:
        return ""

    cleaned = " ".join(date_text.split())
    cleaned = cleaned.replace("–", " - ").replace("—", " - ")

    cleaned = cleaned.replace("Effective number of months achieved:", "")
    cleaned = cleaned.replace("Dates (start-end):", "")
    cleaned = cleaned.strip(" -:")

    return cleaned.strip()


def normalize_single_date(date_str: str) -> str:
    date_str = clean_date_text(date_str)

    patterns = [
        "%d-%m-%Y",
        "%d/%m/%Y",
        "%m-%Y",
        "%m/%Y",
        "%Y-%m-%d",
        "%Y/%m/%d",
    ]

    for pattern in patterns:
        try:
            dt = datetime.strptime(date_str, pattern)
            return dt.strftime("%b %Y")
        except ValueError:
            continue

    return date_str


def parse_date_range(date_text: str | None) -> tuple[str | None, str | None]:
    cleaned = clean_date_text(date_text)
    if not cleaned:
        return None, None

    parts = re.split(r"\s+-\s+", cleaned)

    if len(parts) == 2:
        start_raw = parts[0].strip()
        end_raw = parts[1].strip()
        return normalize_single_date(start_raw), normalize_single_date(end_raw)

    return normalize_single_date(cleaned), None


def parse_projects_from_ec(ec_text: str) -> List[ProjectExperience]:
    projects: List[ProjectExperience] = []

    chunks = ec_text.split("PROJECT EXPERIENCE")
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue

        company = _extract_field(chunk, "Company (employer):")
        project_name = _extract_field(chunk, "Project name:")
        dates = _extract_field(chunk, "Dates (start-end):")
        months = _extract_field(chunk, "Effective number of months achieved:")
        client = _extract_field(chunk, "Client (customer) :") or _extract_field(chunk, "Client (customer):")
        project_size = _extract_field(chunk, "Project size:")
        project_description = _extract_field(chunk, "Project description:")
        role = derive_role(company or "", project_description or "")

        responsibilities_block = _extract_field(
            chunk,
            "External service provider’s roles & responsibilities in the project:",
            stop_markers=["Technologies and methodologies used by the external service provider in the project:"],
        )
        technologies_block = _extract_field(
            chunk,
            "Technologies and methodologies used by the external service provider in the project:",
        )

        if not company and not project_description:
            continue

        responsibilities = split_lines_clean(responsibilities_block)
        technologies = split_technologies(technologies_block)

        start_date, end_date = parse_date_range(dates)

        if not start_date and months:
            start_date, end_date = parse_date_range(months)

        projects.append(
            ProjectExperience(
                company=company or "Unknown",
                project_name=project_name,
                role=role,
                start_date=start_date,
                end_date=end_date,
                months=months,
                client=client,
                project_size=project_size,
                project_description=project_description,
                responsibilities=responsibilities,
                technologies=technologies,
            )
        )

    return projects


def _extract_field(text: str, label: str, stop_markers: list[str] | None = None) -> str | None:
    stop_markers = stop_markers or [
        "Company (employer):",
        "Dates (start-end):",
        "Effective number of months achieved:",
        "Client (customer):",
        "Client (customer) :",
        "Project size:",
        "Project description:",
        "External service provider’s roles & responsibilities in the project:",
        "Technologies and methodologies used by the external service provider in the project:",
        "CV experience page number for this CV:",
    ]

    if label not in text:
        return None

    start = text.find(label) + len(label)
    remaining = text[start:].strip()

    next_positions = []
    for marker in stop_markers:
        pos = remaining.find(marker)
        if pos != -1:
            next_positions.append(pos)

    if next_positions:
        end = min(next_positions)
        return remaining[:end].strip()

    return remaining.strip()


def split_lines_clean(text: str | None) -> List[str]:
    if not text:
        return []

    noise_lines = {
        "External service provider’s roles & responsibilities in the project:",
        "External service provider's roles & responsibilities in the project:",
        "Technologies and methodologies used by the external service provider in the project:",
        "Project description:",
        "Project size:",
        "Client (customer):",
        "Client (customer) :",
    }

    cleaned_lines = []
    seen = set()

    for raw_line in text.splitlines():
        line = raw_line.strip(" -•\t").strip()
        if not line:
            continue

        if line in noise_lines:
            continue

        normalized = " ".join(line.split()).lower()
        if normalized in seen:
            continue

        seen.add(normalized)
        cleaned_lines.append(line)

    return cleaned_lines


def split_technologies(text: str | None) -> List[str]:
    if not text:
        return []

    cleaned = text.replace("Technologies:", "").strip()
    parts = re.split(r",|\n", cleaned)
    return [p.strip() for p in parts if p.strip()]


def build_knowledge_base(master_path: Path, ec_path: Path) -> ResumeKnowledgeBase:
    master_text = read_docx_text(master_path)
    ec_text = read_docx_text(ec_path)

    master_data = parse_master_resume(master_text)

    kb = ResumeKnowledgeBase(
        candidate_profile=master_data["candidate_profile"],
        raw_master_resume_summary=master_data["raw_master_resume_summary"],
        languages=extract_languages_from_ec(ec_text),
        education=extract_education_from_ec(ec_text),
        certifications=extract_certifications_from_ec(ec_text),
        projects=parse_projects_from_ec(ec_text),
    )

    return kb